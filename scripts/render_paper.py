"""
Generic Markdown -> DOCX renderer for this paper's .md source, replacing the
previous approach of hand-mirroring the .md into a separate build_paper.py
with duplicated string literals (which drifted out of sync whenever either
file was edited independently). This renders directly from whichever .md is
given, so the .md is the single source of truth for both the Chinese and
English versions of the manuscript.

Supported markdown subset (matches what this paper's .md actually uses):
  # title                  -> Title (level-0 heading)
  ## heading                -> level-1 heading (h1)
  ### heading                -> level-2 heading (h2)
  **bold** / *italic* inline -> mixed-run paragraphs
  \\* escaped literal asterisk
  ![alt](path)              -> embedded image (path resolved against --figdir)
  a lone *italic line*      -> figure/table caption style (centered, gray, 9pt)
  **Table N. Title** immediately followed by a pipe table -> table + caption
    rendered AFTER the table (caption-below convention used throughout this
    paper), rather than markdown's caption-above order
  | pipe | tables |         -> Word table, auto-sized column widths
  --- (own line)            -> page break
  N. Reference text         -> reference-list paragraph (9.5pt), one per line

Usage: python3 render_paper.py <input.md> <output.docx> [--figdir DIR]
"""
import os
import re
import sys
import argparse

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DEFAULT_FIGDIR = f"{_PROJECT_ROOT}/analysis_output/figures"

ESC_PLACEHOLDER = ""  # private-use codepoint, safe stand-in for an escaped literal "*"


def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.3
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return doc


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


INLINE_TOKEN_RE = re.compile(r'(\*\*.+?\*\*|\*[^\n*]+?\*)', re.S)


def add_inline_runs(paragraph, text, base_size=11, base_bold=False, base_italic=False):
    """Split text on **bold** / *italic* markers and add corresponding runs."""
    text = text.replace('\\*', ESC_PLACEHOLDER)
    pos = 0
    for m in INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], base_size, base_bold, base_italic)
        token = m.group(0)
        if token.startswith('**'):
            _add_run(paragraph, token[2:-2], base_size, True, base_italic)
        else:
            _add_run(paragraph, token[1:-1], base_size, base_bold, True)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], base_size, base_bold, base_italic)


def _add_run(paragraph, text, size, bold, italic):
    text = text.replace(ESC_PLACEHOLDER, '*')
    if text == '':
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def is_whole_line_wrapped(text, marker):
    n = len(marker)
    return text.startswith(marker) and text.endswith(marker) and len(text) > 2 * n


def strip_wrap(text, marker):
    return text[len(marker):-len(marker)]


def parse_table_block(lines):
    rows = []
    for i, line in enumerate(lines):
        if i == 1:
            continue  # separator row (|---|---|)
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows[0], rows[1:]


def compute_col_widths(headers, rows, total_cm=17.0, min_cm=1.6, max_cm=7.0):
    ncols = len(headers)
    maxlen = [len(str(headers[i]))
              for i in range(ncols)]
    for r in rows:
        for i in range(ncols):
            if i < len(r):
                maxlen[i] = max(maxlen[i], len(str(r[i])))
    weights = [max(m, 4) for m in maxlen]
    total_w = sum(weights)
    widths = [max(min_cm, min(max_cm, total_cm * w / total_w)) for w in weights]
    return widths


def render(md_path, out_path, figdir=DEFAULT_FIGDIR):
    with open(md_path, encoding='utf-8') as f:
        raw = f.read()

    blocks = [b for b in raw.split('\n\n')]

    doc = setup_document()

    pending_table_caption = None
    title_done = False
    i = 0
    n = len(blocks)
    while i < n:
        block = blocks[i].strip('\n')
        stripped = block.strip()
        i += 1
        if not stripped:
            continue

        # Title
        if stripped.startswith('# ') and not title_done:
            title_text = stripped[2:].strip()
            title = doc.add_heading(title_text, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_done = True
            continue

        # Headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue

        # Horizontal rule -> page break
        if re.fullmatch(r'-{3,}', stripped):
            doc.add_page_break()
            continue

        # Image
        img_m = re.fullmatch(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_m:
            img_path = img_m.group(2)
            fname = os.path.basename(img_path)
            doc.add_picture(os.path.join(figdir, fname), width=Cm(15.5))
            continue

        # Lone italic line -> caption style
        if is_whole_line_wrapped(stripped, '*') and not stripped.startswith('**'):
            cap_text = strip_wrap(stripped, '*')
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(para, cap_text, base_size=9, base_italic=True)
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            continue

        # Table caption line "**表N. ...**" / "**Table N. ...**" immediately preceding a table
        if is_whole_line_wrapped(stripped, '**'):
            inner = strip_wrap(stripped, '**')
            next_block = blocks[i].strip() if i < n else ''
            if next_block.startswith('|'):
                pending_table_caption = inner
                continue

        # Pipe table
        if stripped.startswith('|'):
            lines = [ln for ln in stripped.split('\n') if ln.strip()]
            headers, rows = parse_table_block(lines)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            for c, h in enumerate(headers):
                hdr_cells[c].text = ''
                p0 = hdr_cells[c].paragraphs[0]
                add_inline_runs(p0, h, base_size=9.5, base_bold=True)
                set_cell_shading(hdr_cells[c], "D9E2F3")
            for r in rows:
                cells = table.add_row().cells
                for c in range(len(headers)):
                    val = r[c] if c < len(r) else ''
                    cells[c].text = ''
                    p0 = cells[c].paragraphs[0]
                    add_inline_runs(p0, val, base_size=9)
            widths = compute_col_widths(headers, rows)
            for row in table.rows:
                for c, w in enumerate(widths):
                    row.cells[c].width = Cm(w)
            doc.add_paragraph()
            if pending_table_caption:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(para, pending_table_caption, base_size=9, base_italic=True)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
                pending_table_caption = None
            continue

        # Reference list (multiple "N. ..." lines inside one block)
        ref_lines = stripped.split('\n')
        if all(re.match(r'^\d+\.\s', ln.strip()) for ln in ref_lines if ln.strip()):
            for ln in ref_lines:
                ln = ln.strip()
                if not ln:
                    continue
                para = doc.add_paragraph()
                add_inline_runs(para, ln, base_size=9.5)
            continue

        # Generic paragraph (mixed bold/italic runs)
        para = doc.add_paragraph()
        add_inline_runs(para, stripped, base_size=11)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('input_md')
    ap.add_argument('output_docx')
    ap.add_argument('--figdir', default=DEFAULT_FIGDIR)
    args = ap.parse_args()
    render(args.input_md, args.output_docx, args.figdir)
